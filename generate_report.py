from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ضبط اتجاه الصفحة والهوامش
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ضبط RTL للمستند كله
def set_rtl_paragraph(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_rtl_table(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

def add_rtl_para(doc, text, style=None, bold=False, size=None, color=None, center=False):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    set_rtl_paragraph(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:hint'), 'cs')
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_heading_styled(doc, text, size, color, bold=True, center=False):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:hint'), 'cs')
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def cell_rtl(cell, text, bold=False, size=11, color=None, bg=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:hint'), 'cs')
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if bg:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)

def shade_row(row, hex_color):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def set_col_width(table, col_index, width_cm):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)

# ===================== COVER PAGE =====================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
set_rtl_paragraph(p)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading_styled(doc, 'تحليل نظام تنجيز', 28, (31, 73, 125), bold=True, center=True)
add_heading_styled(doc, 'منصة تنجيز العقارية — نظام الأوقاف', 20, (68, 114, 196), bold=True, center=True)

doc.add_paragraph()
add_rtl_para(doc, f'تاريخ الإعداد: {datetime.date.today().strftime("%d/%m/%Y")}', bold=False, size=12, center=True)
add_rtl_para(doc, 'إعداد: م.ابراهيم الذارحي', bold=False, size=12, center=True)

doc.add_page_break()

# ===================== 1. مقدمة =====================
add_heading_styled(doc, 'أولاً: مقدمة', 16, (31, 73, 125), bold=True)
add_rtl_para(doc, 'تم إجراء مراجعة تقنية شاملة لنظام الأوقاف (منصة تنجيز العقارية) الذي يتكون من طبقتين رئيسيتين:', size=11)
add_rtl_para(doc, 'الطبقة الأولى: واجهة المستخدم الأمامية (Frontend) المبنية على تقنية ASP.NET Core MVC.', size=11)
add_rtl_para(doc, 'الطبقة الثانية: واجهة برمجة التطبيقات الخلفية (Backend API) المبنية على ASP.NET Core Web API مع قاعدة بيانات SQL Server.', size=11)

doc.add_paragraph()
add_rtl_para(doc, 'الهدف من هذا التقرير هو تحديد الوضع الحالي لكل جزء من النظام بدقة، وتصنيف المكونات إلى:', size=11)
add_rtl_para(doc, '         أجزاء تعمل بشكل كامل وصحيح.', size=11)
add_rtl_para(doc, '         أجزاء متوقفة أو معطّلة بسبب أخطاء في الكود أو روابط خادم منتهية.', size=11)
add_rtl_para(doc, '         أجزاء غير مكتملة بحاجة لمراجعة وإعادة بناء.', size=11)

doc.add_paragraph()

# ===================== 2. الأجزاء الشغالة =====================
add_heading_styled(doc, 'ثانياً: الأجزاء التي تعمل بشكل كامل وصحيح', 16, (31, 73, 125), bold=True)

sections_working = [
    ('تسجيل الدخول وإدارة الجلسة',
     'يعمل نظام تسجيل الدخول بشكل سليم. يُدخل المستخدم اسمه الأول وكلمة المرور، ويقوم النظام بالتحقق من البيانات عبر الـ API وتوليد رمز أمان (JWT Token) وحفظ بيانات الجلسة (نوع المستخدم، رقم المعرف) بشكل صحيح.'),
    ('طلبات التنجيز — الجزء الأساسي من النظام',
     'يعمل هذا القسم بكامله عند الاتصال بالخادم الصحيح. يشمل: عرض قائمة الطلبات، إنشاء طلب جديد، تعديل الطلبات، إدارة حدود الأراضي، والقوائم المنسدلة (أنواع الوثائق، أنواع الملكية، أنواع الأماكن، أغراض المسح).'),
    ('عروض العمل',
     'يعمل هذا القسم بشكل كامل ويشمل: عرض قائمة العروض، إنشاء عرض جديد، تعديله، وتغيير حالته (قبول أو رفض). كذلك يوجد آلية للتحقق من ملف السيرة الذاتية للمستخدم قبل السماح له بالتقديم.'),
    ('إدارة المستخدمين',
     'يعمل هذا القسم ويشمل: تسجيل مستخدم جديد، عرض القائمة الكاملة للمستخدمين، حذف مستخدم، رفع الصورة الشخصية، وعرض التفاصيل الكاملة للمستخدم. كذلك يوجد آلية لتحديث حالة المصادقة.'),
    ('عقود العمل',
     'يعمل هذا القسم ويشمل: إنشاء العقد مع بنوده، عرض قائمة العقود، جلب عقد مهندس معين، واستبدال النصوص في قالب العقد بالقيم الفعلية (مثل التكلفة الإجمالية، التواريخ، البنود). كذلك يُرسل النظام إشعاراً تلقائياً عند إنشاء العقد.'),
    ('الإشعارات',
     'يعمل عرض الإشعارات للمستخدمين بشكل سليم، وتُرسل الإشعارات تلقائياً عند حدوث أحداث مثل إنشاء العقد أو تغيير الحالة.'),
    ('قوالب الرواتب والتنقلات',
     'تعمل هذه الأقسام وتسمح بإنشاء القوالب وحفظ بيانات الجداول المرتبطة.'),
    ('واجهة برمجة التطبيقات الخلفية',
     'تعمل جميع نقاط الوصول (Endpoints) الأساسية في الـ API بشكل صحيح، بما في ذلك: عمليات القراءة والإضافة والتعديل والحذف لجميع الجداول، رفع الصور والملفات، وخدمات الإشعار والتوكن.'),
]

for title, detail in sections_working:
    add_heading_styled(doc, title, 13, (68, 114, 196), bold=True)
    add_rtl_para(doc, detail, size=11)
    doc.add_paragraph()

# ===================== 3. الأجزاء المتوقفة =====================
doc.add_page_break()
add_heading_styled(doc, 'ثالثاً: الأجزاء المتوقفة وأسباب التوقف', 16, (192, 0, 0), bold=True)

doc.add_paragraph()

problems = [
    ('مشكلة روابط الخادم المنتهية — المشكلة الأخطر',
     'يحتوي كنترولر طلبات التنجيز (TanjezOrderController) وكنترولر نموذج العمل (Work_FormController) على روابط ثابتة تشير إلى خادم خارجي قديم هو "tanjeez.somee.com" بدلاً من استخدام الخادم المحلي الفعلي. وكذلك يشير كنترولر الرواتب (WorkWagesTapleController) إلى "awqaf12.somee.com". هذه الروابط لخوادم متوقفة، مما يعني أن أي طلب يُرسَل عبر هذه الصفحات يفشل تماماً رغم أن الخادم المحلي يعمل بشكل صحيح.'),
    ('إشعار "تعليم كمقروء" معطّل',
     'اختيار تعليم الإشعار كمقروء معطّل عن بُعد لأن نقطة الوصول المقابلة في الـ API محاطة بتعليق في الكود (Commented Out) ولا تستجيب لأي طلبات.'),
    ('تحديث حالة العقود معطّل',
     'نقطة الوصول الخاصة بتحديث حالة العقد (UpdateStatus) موجودة في الكود لكنها غير فعّالة لأنها علّقها المبرمج، مما يعني أن تغيير حالة العقد لن يُحفظ في قاعدة البيانات.'),
    ('حماية الصلاحيات غير مكتملة',
     'لا توجد حماية على مستوى واجهة الـ API. جميع نقاط الوصول مفتوحة لأي طلب دون التحقق من هوية المستخدم أو صلاحياته. في الواجهة الأمامية، بعض الصفحات تتحقق من الجلسة يدوياً لكن ذلك غير موحّد، إذ يمكن الوصول لكثير من الصفحات دون تسجيل الدخول.'),
    ('خطأ منطقي في التحقق من السيرة الذاتية',
     'كود التحقق من وجود سيرة ذاتية للمستخدم يحتوي على خطأ منطقي يجعل الشرط صحيحاً دائماً، مما يعني أن النظام سيطلب رفع سيرة ذاتية من كل مستخدم في كل مرة بغض النظر عن وضعه الفعلي.'),
    ('صفحات موجودة بدون منطق',
     'توجد عدة صفحات في الواجهة الأمامية مبنية بصرياً لكن بدون منطق وظيفي خلفها، مثل: صفحات الخدمات، صفحة الأعمال (Works)، وبعض صفحات الحساب. هذه الصفحات تظهر للمستخدم لكنها لا تؤدي أي وظيفة.'),
]

for title, detail in problems:
    add_heading_styled(doc, title, 13, (192, 0, 0), bold=True)
    add_rtl_para(doc, detail, size=11)
    doc.add_paragraph()

# ===================== 4. ملاحظات المبرمج السابق =====================
doc.add_page_break()
add_heading_styled(doc, 'رابعاً: ملاحظات على أسلوب العمل السابق', 16, (131, 60, 12), bold=True)
doc.add_paragraph()

obs = [
    ('نسخ متعددة من نفس الصفحة',
     'تركت عملية التطوير وراءها نسخاً متعددة من بعض الصفحات دون أي تنسيق أو تنظيف. على سبيل المثال، توجد سبع نسخ من صفحة إنشاء نموذج النقل (Create, Create1, Create2, Create3, Create4, Create5, Create6) مما يدل على تجريب متكرر دون حذف الإصدارات القديمة. الوضع مشابه في نماذج العمل وأنواع العروض.'),
    ('كود معلّق بكميات كبيرة',
     'يحتوي الكود على كميات كبيرة من الكود المعلّق (Commented Out) في معظم الملفات، مما يصعّب فهم ما هو شغّال وما هو غير شغّال دون فحص دقيق. هذا يشير إلى أن المبرمج كان يبدّل بين تجارب مختلفة دون تنظيف.'),
    ('مشكلة ترميز البيانات العربية',
     'البيانات الأولية التي أُدخلت عند أول تشغيل للنظام كانت تالفة بسبب خطأ في حفظ ملفات الكود بترميز غير صحيح، مما أدى إلى تخزين نصوص عربية مشوّهة في قاعدة البيانات. تم اكتشاف هذه المشكلة ومعالجتها في المرحلة الحالية.'),
]

for title, detail in obs:
    add_heading_styled(doc, title, 13, (131, 60, 12), bold=True)
    add_rtl_para(doc, detail, size=11)
    doc.add_paragraph()

# ===================== 5. جدول ملخص =====================
doc.add_page_break()
add_heading_styled(doc, 'خامساً: جدول ملخص الوضع الحالي', 16, (31, 73, 125), bold=True)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
set_rtl_table(table)

hdr = table.rows[0]
shade_row(hdr, '1F497D')
cell_rtl(hdr.cells[2], 'المكوّن', bold=True, size=11, color=(255,255,255))
cell_rtl(hdr.cells[1], 'الوضع', bold=True, size=11, color=(255,255,255))
cell_rtl(hdr.cells[0], 'التفاصيل', bold=True, size=11, color=(255,255,255))

rows_data = [
    ('تسجيل الدخول', 'يعمل', 'يعمل بشكل كامل وصحيح'),
    ('طلبات التنجيز', 'خلل في الرابط', 'يشير لخادم somee.com متوقف بدلاً من الخادم المحلي'),
    ('عروض العمل', 'يعمل', 'يعمل بشكل كامل وصحيح'),
    ('عقود العمل', 'يعمل جزئياً', 'الإنشاء يعمل، تحديث الحالة معطّل'),
    ('نظام الإشعارات', 'يعمل جزئياً', 'العرض يعمل، تعليم كمقروء معطّل'),
    ('إدارة المستخدمين', 'يعمل', 'يعمل بشكل كامل وصحيح'),
    ('قوالب الرواتب', 'يعمل', 'يعمل بشكل كامل وصحيح'),
    ('نموذج العمل (Work Form)', 'خلل في الرابط', 'يشير لخادم somee.com متوقف'),
    ('حماية الصلاحيات', 'ناقص', 'الـ API مفتوح بالكامل بدون حماية'),
    ('بيانات قاعدة البيانات', 'تم إصلاحها', 'كانت البيانات العربية تالفة وتم إصلاحها'),
    ('صفحات الخدمات والأعمال', 'شكلي فقط', 'صفحات موجودة بدون وظيفة فعلية'),
    ('التحقق من السيرة الذاتية', 'خطأ منطقي', 'الشرط خاطئ ويطلب الرفع دائماً'),
]

colors_map = {
    'يعمل': 'C6EFCE',
    'يعمل جزئياً': 'FFEB9C',
    'خلل في الرابط': 'FFC7CE',
    'ناقص': 'FFC7CE',
    'تم إصلاحها': 'C6EFCE',
    'شكلي فقط': 'FFEB9C',
    'خطأ منطقي': 'FFC7CE',
}
text_colors_map = {
    'يعمل': (0,97,0),
    'يعمل جزئياً': (156,87,0),
    'خلل في الرابط': (156,0,6),
    'ناقص': (156,0,6),
    'تم إصلاحها': (0,97,0),
    'شكلي فقط': (156,87,0),
    'خطأ منطقي': (156,0,6),
}

for comp, status, detail in rows_data:
    row = table.add_row()
    bg = colors_map.get(status, 'FFFFFF')
    tc = text_colors_map.get(status, (0,0,0))
    cell_rtl(row.cells[2], comp, bold=True, size=11)
    cell_rtl(row.cells[1], status, bold=True, size=11, color=tc, bg=bg)
    cell_rtl(row.cells[0], detail, size=10)

set_col_width(table, 2, 4.5)
set_col_width(table, 1, 3.5)
set_col_width(table, 0, 7.5)

# ===================== 6. التوصيات =====================
doc.add_page_break()
add_heading_styled(doc, 'سادساً: التوصيات وأولويات الإصلاح', 16, (31, 73, 125), bold=True)
doc.add_paragraph()

recs = [
    ('الأولوية الأولى — تصحيح روابط الخادم',
     'يجب على الفور استبدال الروابط الثابتة القديمة (somee.com) في ثلاثة كنترولرات بالرابط الصحيح المُعرَّف في ملف الإعدادات. هذا الإصلاح سيعيد تشغيل القسم الأساسي من النظام وهو طلبات التنجيز ونماذج العمل.'),
    ('الأولوية الثانية — تفعيل المميزات المعلّقة',
     'يجب تفعيل نقطتي الوصول المعلّقتين في الـ API وهما: تعليم الإشعار كمقروء، وتحديث حالة العقد. هذا يتطلب إزالة تعليق بضعة أسطر في الكود.'),
    ('الأولوية الثالثة — إضافة حماية الصلاحيات',
     'يجب إضافة آلية موحّدة للتحقق من هوية المستخدم على مستوى الـ API، وتحديد الصلاحيات بحسب نوع المستخدم (عميل أو مقاول أو مدير).'),
    ('الأولوية الرابعة — تنظيف الكود',
     'يُنصح بحذف الملفات والصفحات والكود المعلّق غير المستخدم، وتوحيد أسلوب العمل. هذا سيسهّل الصيانة المستقبلية ويقلل من احتمالية الأخطاء.'),
    ('الأولوية الخامسة — إصلاح الأخطاء المنطقية',
     'يجب مراجعة وإصلاح الشرط الخاطئ في آلية التحقق من السيرة الذاتية، وتصحيح القيم المكررة في قالب العقد.'),
]

for i, (title, detail) in enumerate(recs, 1):
    add_heading_styled(doc, f'التوصية {i}: {title}', 13, (31, 73, 125), bold=True)
    add_rtl_para(doc, detail, size=11)
    doc.add_paragraph()

# ===================== الخاتمة =====================
add_heading_styled(doc, 'خاتمة', 14, (31, 73, 125), bold=True)
add_rtl_para(doc,
    'النظام في مجمله يمتلك بنية تقنية جيدة وقابلة للتطوير. المشاكل الموجودة قابلة للحل وليست بنيوية، '
    'بل تتركز في أخطاء إعداد الروابط وعدم اكتمال بعض المميزات. '
    'بعد تطبيق التوصيات المذكورة، يُتوقع أن يعمل النظام بكامل وظائفه المخطط لها.',
    size=11)

doc.add_paragraph()
add_rtl_para(doc, f'انتهى التقرير — {datetime.date.today().strftime("%d/%m/%Y")}', size=10, center=True)

output_path = r'd:\alnahari\تقرير_تحليل_نظام_الأوقاف.docx'
doc.save(output_path)
import sys
sys.stdout.buffer.write(b'SUCCESS: Report generated\n')
